import os
import requests
import base64
import docx2txt
import tempfile
from pypdf import PdfReader
from docx import Document
from Xfrate2.utils import logger
from Xfrate2.state import AgentState

# def parse_document(state: AgentState) -> dict:
#     """
#     Node 1: Reads the file from disk and extracts text or encodes images.
#     """
#     file_path = state.get("file_path")
    
#     if not file_path or not os.path.exists(file_path):
#         logger.error(f"File not found: {file_path}")
#         raise FileNotFoundError(f"File not found: {file_path}")

#     logger.info(f"START: Processing file: {file_path}")
    
#     _, ext = os.path.splitext(file_path)
#     ext = ext.lower()
    
#     extracted_text = ""
#     file_type = ext
    
#     try:
#         # --- CASE A: PDF FILES ---
#         if ext == ".pdf":
#             reader = PdfReader(file_path)
#             text_content = []
#             for page in reader.pages:
#                 text_content.append(page.extract_text())
#             extracted_text = "\n".join(text_content)
#             logger.info(f"Extracted {len(extracted_text)} characters from PDF.")

#         # --- CASE B: WORD DOCS ---
#         elif ext == ".docx":
#             # doc = Document(file_path)
#             doc =docx2txt.process(file_path)
#             full_text = []
            
#             # 1. Read Paragraphs (Standard text)
#             # for para in doc.paragraphs:
#             #     full_text.append(para.text)
            
#             # # 2. Read Tables (CRITICAL FIX)
#             # for table in doc.tables:
#             #     for row in table.rows:
#             #         # Join cells with a separator (pipe |) so LLM understands row structure
#             #         row_text = [cell.text for cell in row.cells]
#             #         full_text.append(" | ".join(row_text))
            
#             extracted_text =doc
#             logger.info(f"Extracted {len(extracted_text)} characters from DOCX.")

#         # --- CASE C: IMAGES (OCR needed by LLM) ---
#         elif ext in [".png", ".jpg", ".jpeg"]:
#             # file_type = True
#             with open(file_path, "rb") as image_file:
#                 # We store the base64 string as the "text" payload for the Vision model
#                 extracted_text = base64.b64encode(image_file.read()).decode('utf-8')
#             logger.info(f"Encoded image to base64 for Vision model.")

#         # --- CASE D: TEXT FILES ---
#         elif ext == ".txt":
#             with open(file_path, "r", encoding="utf-8") as f:
#                 extracted_text = f.read()
#             logger.info(f"Read {len(extracted_text)} characters from text file.")

#         else:
#             raise ValueError(f"Unsupported file format: {ext}")

#         # Update the State
#         return {
#             "extracted_text": extracted_text,
#             "file_type": file_type
#         }

#     except Exception as e:
#         logger.error(f"Failed to parse document: {e}", exc_info=True)
#         raise e
    

def parse_document(state: AgentState) -> dict:
    """
    Node 1 (API Version): 
    1. Downloads file from state['document_url'].
    2. Runs standard text extraction (PyPDF, Docx, etc.).
    """
    doc_url = state.get("document_url")
    if not doc_url:
        raise ValueError("Missing 'document_url' in state.")

    logger.info(f"Downloading document from: {doc_url}")

    # 1. Download File
    try:
        response = requests.get(doc_url, stream=True)
        response.raise_for_status()
        
        # Infer extension from URL (fallback to .pdf)
        filename = doc_url.split("?")[0].split("/")[-1]
        _, ext = os.path.splitext(filename)
        if not ext:
            ext = ".pdf"
        ext = ext.lower()

        # Create Temp File
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
            for chunk in response.iter_content(chunk_size=8192):
                tmp_file.write(chunk)
            temp_path = tmp_file.name
            
        logger.info(f"File saved to temp: {temp_path}")

    except Exception as e:
        logger.error(f"Download Failed: {e}")
        raise e

    # 2. Extract Content (Your Existing Logic)
    extracted_text = ""
    try:
        # --- CASE A: PDF FILES ---
        if ext == ".pdf":
            reader = PdfReader(temp_path)
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text() or "")
            extracted_text = "\n".join(text_content)

        # --- CASE B: WORD DOCS ---
        elif ext == ".docx":
            doc = Document(temp_path)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])

        # --- CASE C: IMAGES ---
        elif ext in [".png", ".jpg", ".jpeg"]:
            with open(temp_path, "rb") as image_file:
                extracted_text = base64.b64encode(image_file.read()).decode('utf-8')

        # --- CASE D: TEXT FILES ---
        elif ext == ".txt":
            with open(temp_path, "r", encoding="utf-8") as f:
                extracted_text = f.read()
        
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        logger.info(f"Extraction complete. {len(extracted_text)} chars.")

    finally:
        # Cleanup: Delete the temp file to keep server clean
        if os.path.exists(temp_path):
            os.remove(temp_path)
            logger.info("Temp file cleaned up.")

    # Return updated state
    return {
        "extracted_text": extracted_text,
        "file_type": ext,
        "file_path": doc_url # Keep the URL as the source of truth for metadata
    }